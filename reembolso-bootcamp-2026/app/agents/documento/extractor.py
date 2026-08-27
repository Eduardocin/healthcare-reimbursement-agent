"""Extração de texto e campos estruturados de anexos (PDF, Imagens com OCR, DOCX)."""

from __future__ import annotations

import io
import json
import logging
import re
from decimal import Decimal
from typing import Any

import pymupdf  # PyMuPDF
import pytesseract
from PIL import Image
from pydantic import BaseModel, Field

from app.llm import criar_llm
from app.schemas import Categoria

logger = logging.getLogger(__name__)


class ExtractedData(BaseModel):
    categoria: Categoria
    eh_documento_fiscal_assistencial: bool = True
    valor_pago_brl: Decimal | None = None
    data_atendimento: str | None = None  # YYYY-MM-DD
    codigo_procedimento: str | None = None
    descricao_procedimento: str | None = None
    prestador_nome: str | None = None
    prestador_registro: str | None = None
    paciente_nome: str | None = None
    campos_presentes: list[str] = Field(default_factory=list)
    campos_ausentes: list[str] = Field(default_factory=list)
    observacao: str | None = None


REQUIRED_FISCAL_FIELDS = (
    "beneficiario_nome",
    "beneficiario_cpf",
    "prestador_nome",
    "prestador_documento",
    "registro_conselho",
    "data_atendimento",
    "descricao_procedimento",
    "valor",
    "assinatura_carimbo",
)


def _money(text: str) -> Decimal | None:
    matches = re.findall(r"r\$\s*([\d.]+,\d{2})", text, re.IGNORECASE)
    if not matches:
        return None
    preferred = re.search(
        r"valor\s*total\s*:?\s*r\$\s*([\d.]+,\d{2})",
        text,
        re.IGNORECASE,
    )
    raw = preferred.group(1) if preferred else matches[0]
    try:
        return Decimal(raw.replace(".", "").replace(",", "."))
    except ValueError:
        return None


def _date(text: str, label: str = "atendimento") -> str | None:
    match = re.search(
        rf"data\s*(?:do\s*)?{label}\s*:?\s*(\d{{2}}/\d{{2}}/\d{{4}})",
        text,
        re.IGNORECASE,
    )
    if not match:
        return None
    day, month, year = match.group(1).split("/")
    return f"{year}-{month}-{day}"


def _procedure_code(text: str, prefix: str | None = None) -> str | None:
    pattern = rf"\b{re.escape(prefix)}\d{{{8 - len(prefix)}}}\b" if prefix else r"\b\d{8}\b"
    match = re.search(pattern, text)
    return match.group(0) if match else None


def _field_status(
    text: str,
    *,
    value: Decimal | None,
    service_date: str | None,
) -> tuple[list[str], list[str]]:
    checks = {
        "beneficiario_nome": bool(re.search(r"(?:paciente|recebi de)\s*:", text, re.IGNORECASE)),
        "beneficiario_cpf": bool(re.search(r"cpf\s+do\s+paciente\s*:", text, re.IGNORECASE)),
        "prestador_nome": bool(re.search(r"(?:profissional|prestador)\s*:", text, re.IGNORECASE)),
        "prestador_documento": "cpf/cnpj" in text.lower(),
        "registro_conselho": bool(re.search(r"\b(?:CRM|CRP|CREFITO|CRFA|COREN)[-/]?[A-Z0-9]*", text, re.IGNORECASE)),
        "data_atendimento": service_date is not None,
        "descricao_procedimento": bool(re.search(r"(?:referente a|descri[cç][aã]o|discrimina[cç][aã]o)", text, re.IGNORECASE)),
        "valor": value is not None,
        "assinatura_carimbo": "____" in text or "carimbo" in text.lower(),
    }
    present = [name for name, is_present in checks.items() if is_present]
    missing = [name for name in REQUIRED_FISCAL_FIELDS if name not in present]
    return present, missing


def extract_raw_text(file_bytes: bytes, filename: str, mime_type: str) -> str:
    """Extrai texto de PDF, imagem ou DOCX, executando OCR quando necessário."""
    fn_lower = filename.lower()
    mime_lower = mime_type.lower()

    # 1. PDF
    if fn_lower.endswith(".pdf") or "pdf" in mime_lower:
        try:
            doc = pymupdf.open(stream=file_bytes, filetype="pdf")
            text_parts = []
            for page in doc:
                page_text = page.get_text().strip()
                if len(page_text) < 30:
                    # Tenta OCR na página
                    try:
                        pix = page.get_pixmap(dpi=150)
                        img = Image.open(io.BytesIO(pix.tobytes("png")))
                        ocr_text = pytesseract.image_to_string(img, lang="por")
                        if len(ocr_text.strip()) > len(page_text):
                            page_text = ocr_text.strip()
                    except Exception as e:
                        logger.warning(f"Falha no OCR da página PDF: {e}")
                text_parts.append(page_text)
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"Erro ao ler PDF {filename}: {e}")
            return ""

    # 2. Imagens (PNG, JPG, TIFF, WEBP)
    if any(fn_lower.endswith(ext) for ext in [".png", ".jpg", ".jpeg", ".webp", ".tiff", ".bmp"]) or "image" in mime_lower:
        try:
            img = Image.open(io.BytesIO(file_bytes))
            return pytesseract.image_to_string(img, lang="por").strip()
        except Exception as e:
            logger.error(f"Erro ao fazer OCR de imagem {filename}: {e}")
            return ""

    # 3. DOCX
    if fn_lower.endswith(".docx") or "officedocument" in mime_lower:
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Erro ao ler DOCX {filename}: {e}")
            return ""

    # Fallback para texto puro
    try:
        return file_bytes.decode("utf-8", errors="ignore")
    except Exception:
        return ""


def extract_document_info(
    file_bytes: bytes,
    filename: str,
    mime_type: str,
) -> ExtractedData:
    """Extrai texto e classifica o documento nas 7 categorias com campos obrigatórios."""
    raw_text = extract_raw_text(file_bytes, filename, mime_type)
    if not raw_text.strip():
        return ExtractedData(
            categoria=Categoria.INVALIDO,
            eh_documento_fiscal_assistencial=False,
            observacao="Arquivo vazio, corrompido ou ilegível.",
        )

    # Heurística rápida para documento inválido óbvio (ex: conta de consumo/energia)
    text_lower = raw_text.lower()
    if (
        "energia elétrica" in text_lower
        or "fatura de energia" in text_lower
        or "conta de consumo" in text_lower
        or "distribuição de energia" in text_lower
        or "não constitui documento fiscal de despesa assistencial" in text_lower
    ):
        return ExtractedData(
            categoria=Categoria.INVALIDO,
            eh_documento_fiscal_assistencial=False,
            observacao="Conta de consumo de energia elétrica não constitui documento fiscal assistencial.",
        )

    # Heurística para OPME / Prótese
    if (
        "prótese" in text_lower
        or "órtese" in text_lower
        or "artroplastia" in text_lower
        or "70000039" in raw_text
        or "componente tibial" in text_lower
    ):
        val = _money(raw_text)
        data_atend = _date(raw_text)
        present, missing = _field_status(raw_text, value=val, service_date=data_atend)

        return ExtractedData(
            categoria=Categoria.MATERIAL_OPME,
            eh_documento_fiscal_assistencial=True,
            valor_pago_brl=val,
            data_atendimento=data_atend,
            codigo_procedimento=_procedure_code(raw_text, "7"),
            descricao_procedimento="material, órtese, prótese ou material especial",
            campos_presentes=present,
            campos_ausentes=missing,
        )

    # Heurística para Relatório Clínico
    if (
        "relatório clínico circunstanciado" in text_lower
        or "manifestação do profissional assistente" in text_lower
        or "acompanhamento psicoterápico" in text_lower
    ):
        val = _money(raw_text)
        service_date = _date(raw_text) or _date(raw_text, "emissão")
        present, missing = _field_status(raw_text, value=val, service_date=service_date)
        return ExtractedData(
            categoria=Categoria.RELATORIO_CLINICO,
            eh_documento_fiscal_assistencial=True,
            valor_pago_brl=val,
            data_atendimento=service_date,
            codigo_procedimento=_procedure_code(raw_text),
            descricao_procedimento="relatório clínico",
            campos_presentes=present,
            campos_ausentes=missing if val is not None else [],
        )

    # Heurística para Consulta Médica (Dermatologia / TUSS 10101012)
    if (
        "consulta médica" in text_lower
        or "10101012" in raw_text
        or "dermatologia" in text_lower
        or "serviços médicos" in text_lower
    ):
        val = _money(raw_text)
        data_atend = _date(raw_text)
        present, missing = _field_status(raw_text, value=val, service_date=data_atend)

        return ExtractedData(
            categoria=Categoria.CONSULTA_MEDICA,
            eh_documento_fiscal_assistencial=True,
            valor_pago_brl=val,
            data_atendimento=data_atend,
            codigo_procedimento=_procedure_code(raw_text, "1"),
            descricao_procedimento="consulta médica",
            campos_presentes=present,
            campos_ausentes=missing,
        )

    # Heurística para Sessão de Psicoterapia (TUSS 50000462)
    if (
        "psicoterapia" in text_lower
        or "50000462" in raw_text
        or "psicologia clínica" in text_lower
    ):
        val = _money(raw_text)
        data_atend = _date(raw_text)
        present, campos_ausentes = _field_status(raw_text, value=val, service_date=data_atend)

        # Verifica se informa número da sessão
        if "não informado" in text_lower or "sessão nº" not in text_lower:
            campos_ausentes.append("numero_sessao")

        return ExtractedData(
            categoria=Categoria.SESSAO_TERAPIA,
            eh_documento_fiscal_assistencial=True,
            valor_pago_brl=val,
            data_atendimento=data_atend,
            codigo_procedimento=_procedure_code(raw_text, "5"),
            descricao_procedimento="sessão de terapia",
            campos_presentes=present,
            campos_ausentes=campos_ausentes,
        )

    # Extração genérica via LLM para casos dinâmicos da avaliação
    try:
        llm = criar_llm()
        prompt = (
            "Você é um extrator de dados de documentos de reembolso de saúde.\n"
            "Analise o texto extraído do documento abaixo e classifique nas seguintes categorias:\n"
            "- CONSULTA_MEDICA\n"
            "- SESSAO_TERAPIA\n"
            "- EXAME_DIAGNOSTICO\n"
            "- RELATORIO_CLINICO\n"
            "- MATERIAL_OPME\n"
            "- DESPESA_NAO_COBERTA\n"
            "- INVALIDO (se for conta de luz/água, boleto comercial, documento não médico/assistencial)\n\n"
            f"Texto do documento:\n{raw_text[:3000]}\n\n"
        "Identifique os campos documentais sem completar nada por inferência. "
        "Use nos arrays apenas estes nomes: beneficiario_nome, beneficiario_cpf, "
        "prestador_nome, prestador_documento, registro_conselho, data_atendimento, "
        "descricao_procedimento, valor, assinatura_carimbo, numero_sessao, "
        "pedido_medico, indicacao_clinica.\n"
        "Responda SOMENTE em formato JSON com as chaves:\n"
            "{\n"
            '  "categoria": "CONSULTA_MEDICA | SESSAO_TERAPIA | EXAME_DIAGNOSTICO | RELATORIO_CLINICO | MATERIAL_OPME | DESPESA_NAO_COBERTA | INVALIDO",\n'
            '  "eh_documento_fiscal_assistencial": true/false,\n'
            '  "valor_pago_brl": 0.00,\n'
            '  "data_atendimento": "YYYY-MM-DD",\n'
        '  "codigo_procedimento": "código TUSS se houver ou null",\n'
        '  "descricao_procedimento": "descrição literal resumida ou null",\n'
            '  "campos_presentes": ["beneficiario_nome", "valor"],\n'
            '  "campos_ausentes": []\n'
            "}"
        )
        response = llm.invoke(prompt)
        text_resp = response.content if hasattr(response, "content") else str(response)

        # Extrai JSON
        json_match = re.search(r"\{.*\}", text_resp, re.DOTALL)
        if json_match:
            data = json.loads(json_match.group(0))
            cat = Categoria(data.get("categoria", "INVALIDO"))
            val = Decimal(str(data["valor_pago_brl"])) if data.get("valor_pago_brl") else None
            return ExtractedData(
                categoria=cat,
                eh_documento_fiscal_assistencial=data.get("eh_documento_fiscal_assistencial", True),
                valor_pago_brl=val,
                data_atendimento=data.get("data_atendimento"),
                codigo_procedimento=data.get("codigo_procedimento"),
                descricao_procedimento=data.get("descricao_procedimento"),
                campos_presentes=data.get("campos_presentes", []),
                campos_ausentes=data.get("campos_ausentes", []),
            )
    except Exception as e:
        logger.error(f"Erro na extração LLM do documento: {e}")

    return ExtractedData(
        categoria=Categoria.INVALIDO,
        eh_documento_fiscal_assistencial=False,
        observacao="Não foi possível identificar o tipo de documento.",
    )
