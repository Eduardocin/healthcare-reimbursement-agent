"""Extração e chunking determinísticos da base normativa."""

from __future__ import annotations

import hashlib
import re
import unicodedata
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Iterator, cast

import pymupdf
from docx import Document as DocxDocument
from llama_index.core import Document
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core.schema import BaseNode

TUSS_LINE = re.compile(r"(?m)^(?P<code>\d{8})\s*$")
ARTICLE_START = re.compile(r"(?m)(?=^Art\.\s*\d+[º.]?)")
ARTICLE_REFERENCE = re.compile(r"\bArt\.\s*(\d+)", re.IGNORECASE)
TUSS_REFERENCE = re.compile(r"\b(?:TUSS[-\s]*)?(\d{8})\b", re.IGNORECASE)
CIRC_REFERENCE = re.compile(
    r"\bCIRC(?:ULAR(?:\s+NORMATIVA)?)?[-\s]*(\d{1,2})[/-](\d{4})\b",
    re.IGNORECASE,
)

MONTHS = {
    "janeiro": 1,
    "fevereiro": 2,
    "marco": 3,
    "abril": 4,
    "maio": 5,
    "junho": 6,
    "julho": 7,
    "agosto": 8,
    "setembro": 9,
    "outubro": 10,
    "novembro": 11,
    "dezembro": 12,
}


@dataclass(frozen=True)
class SourceUnit:
    text: str
    source: str
    page: int | None
    unit: int
    metadata: dict[str, str | int | float]


def _without_accents(value: str) -> str:
    normalized = unicodedata.normalize("NFKD", value)
    return "".join(char for char in normalized if not unicodedata.combining(char))


def _clean_text(text: str) -> str:
    lines: list[str] = []
    for raw_line in text.replace("\u00a0", " ").splitlines():
        line = " ".join(raw_line.split())
        if not line:
            lines.append("")
            continue
        if re.fullmatch(r"pág\.\s*\d+", line, re.IGNORECASE):
            continue
        if "SaúdeMais Saúde Suplementar S.A. — documento normativo" in line:
            continue
        lines.append(line)
    cleaned = "\n".join(lines)
    return re.sub(r"\n{3,}", "\n\n", cleaned).strip()


def _document_type(path: Path) -> str:
    name = path.stem
    if name == "regulamento_geral":
        return "regulamento"
    if name.startswith("circular_"):
        return "circular"
    if name == "tabela_urs_2026":
        return "tabela_urs"
    if name.startswith("nota_tecnica"):
        return "nota_tecnica"
    if name.startswith("anexo_"):
        return "anexo"
    if name == "faq_interno":
        return "material_apoio"
    return "manual_apoio"


def _authority_rank(document_type: str) -> int:
    return 3 if document_type in {
        "regulamento", "circular", "tabela_urs", "nota_tecnica", "anexo"
    } else 1


def _source_rule_id(path: Path) -> str | None:
    match = re.fullmatch(r"circular_(\d{2})_(\d{4})", path.stem)
    if match:
        return f"CIRC-{int(match.group(1)):02d}-{match.group(2)}"
    if path.stem == "nota_tecnica_02_documentos":
        return "NT-02"
    if path.stem == "anexo_iv_exclusoes":
        return "ANEXO-IV"
    return None


def _effective_date(text: str, path: Path) -> str:
    normalized = _without_accents(text.lower())
    match = re.search(
        r"inicio de vigencia:\s*(\d{1,2})\D{0,3}de ([a-z]+) de (\d{4})",
        normalized,
    )
    if match and match.group(2) in MONTHS:
        day = int(match.group(1))
        month = MONTHS[match.group(2)]
        return f"{match.group(3)}-{month:02d}-{day:02d}"
    if path.stem in {"regulamento_geral", "tabela_urs_2026"}:
        return "2026-01-01"
    return ""


def extract_rule_ids(
    text: str,
    source_rule_id: str | None = None,
    *,
    include_article_ids: bool = True,
) -> list[str]:
    identifiers: set[str] = set()
    if source_rule_id:
        identifiers.add(source_rule_id)
    if include_article_ids:
        identifiers.update(
            f"ART-{int(number)}" for number in ARTICLE_REFERENCE.findall(text)
        )
    identifiers.update(f"TUSS-{code}" for code in TUSS_REFERENCE.findall(text))
    identifiers.update(
        f"CIRC-{int(number):02d}-{year}"
        for number, year in CIRC_REFERENCE.findall(text)
    )
    return sorted(identifiers)


def _base_metadata(path: Path, effective_date: str) -> dict[str, str | int | float]:
    document_type = _document_type(path)
    return {
        "source": path.name,
        "document_type": document_type,
        "authority_rank": _authority_rank(document_type),
        "effective_date": effective_date,
    }


def _split_articles(text: str) -> list[str]:
    starts = [match.start() for match in ARTICLE_START.finditer(text)]
    if not starts:
        return [text]
    units: list[str] = []
    if text[: starts[0]].strip():
        units.append(text[: starts[0]].strip())
    for position, start in enumerate(starts):
        end = starts[position + 1] if position + 1 < len(starts) else len(text)
        units.append(text[start:end].strip())
    return [unit for unit in units if unit]


def _table_units(text: str) -> list[tuple[str, str | None]]:
    matches = list(TUSS_LINE.finditer(text))
    if not matches:
        return [(text, None)]
    units: list[tuple[str, str | None]] = []
    preamble = text[: matches[0].start()].strip()
    if preamble:
        units.append((preamble, None))
    for position, match in enumerate(matches):
        end = matches[position + 1].start() if position + 1 < len(matches) else len(text)
        units.append((text[match.start():end].strip(), match.group("code")))
    return units


def _pdf_units(path: Path) -> Iterator[SourceUnit]:
    document = pymupdf.open(path)
    full_text = "\n".join(
        cast(str, document[index].get_text("text"))
        for index in range(len(document))
    )
    effective_date = _effective_date(full_text, path)
    base = _base_metadata(path, effective_date)
    source_rule_id = _source_rule_id(path)
    unit_index = 0

    for page_index in range(len(document)):
        page_text = _clean_text(cast(str, document[page_index].get_text("text")))
        if not page_text:
            continue
        if path.stem == "tabela_urs_2026":
            raw_units = _table_units(page_text)
        else:
            raw_units = [(unit, None) for unit in _split_articles(page_text)]
        for text, procedure_code in raw_units:
            metadata = {**base, "page": page_index + 1}
            if procedure_code:
                metadata["procedure_code"] = procedure_code
            if source_rule_id:
                metadata["source_rule_id"] = source_rule_id
            yield SourceUnit(text, path.name, page_index + 1, unit_index, metadata)
            unit_index += 1


def _docx_units(path: Path) -> Iterator[SourceUnit]:
    document = DocxDocument(str(path))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                paragraphs.append(" | ".join(cells))
    text = _clean_text("\n\n".join(paragraphs))
    metadata = _base_metadata(path, _effective_date(text, path))
    yield SourceUnit(text, path.name, None, 0, metadata)


def extract_units(path: Path) -> Iterable[SourceUnit]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _pdf_units(path)
    if suffix == ".docx":
        return _docx_units(path)
    raise ValueError(f"formato não suportado na base: {path.name}")


def build_nodes(
    kb_dir: Path,
    chunk_size: int = 512,
    chunk_overlap: int = 80,
) -> list[BaseNode]:
    """Extrai a KB e devolve nós LlamaIndex com IDs estáveis entre builds."""
    splitter = SentenceSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    nodes: list[BaseNode] = []

    for path in sorted(kb_dir.iterdir(), key=lambda item: item.name):
        if path.suffix.lower() not in {".pdf", ".docx"}:
            continue
        for unit in extract_units(path):
            document_id = hashlib.sha256(
                f"{unit.source}:{unit.page}:{unit.unit}".encode("utf-8")
            ).hexdigest()
            document = Document(text=unit.text, metadata=unit.metadata, id_=document_id)
            unit_nodes = splitter.get_nodes_from_documents([document], show_progress=False)
            for chunk_index, node in enumerate(unit_nodes):
                rule_ids = extract_rule_ids(
                    node.get_content(),
                    str(unit.metadata.get("source_rule_id") or "") or None,
                    include_article_ids=unit.metadata["document_type"] != "circular",
                )
                node.metadata["rule_ids"] = ",".join(rule_ids)
                stable_key = (
                    f"{unit.source}:{unit.page}:{unit.unit}:{chunk_index}:"
                    f"{hashlib.sha256(node.get_content().encode('utf-8')).hexdigest()}"
                )
                node.id_ = hashlib.sha256(stable_key.encode("utf-8")).hexdigest()
                nodes.append(node)

    if not nodes:
        raise ValueError(f"nenhum documento indexável encontrado em {kb_dir}")
    return nodes


def source_hashes(kb_dir: Path) -> dict[str, str]:
    return {
        path.name: hashlib.sha256(path.read_bytes()).hexdigest()
        for path in sorted(kb_dir.iterdir(), key=lambda item: item.name)
        if path.suffix.lower() in {".pdf", ".docx"}
    }
